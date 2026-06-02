"""Minimal Markdown -> DOCX renderer.

Handles the subset of Markdown the document generator emits: ATX headings, pipe
tables (with or without a `---` separator row), bullet/numbered lists, fenced code
blocks (including ```mermaid), horizontal rules, and inline **bold** / `code`.

The point of DOCX output is that pipe tables render as real Word tables instead of
collapsing to one line (the rendering complaint on the PDF/DOCX exports).
"""
from __future__ import annotations

import re
from io import BytesIO
from typing import List

_INLINE_RE = re.compile(r"(\*\*.+?\*\*|`[^`]+`)")
_HEADING_RE = re.compile(r"^(#{1,6})\s+(.*)$")
_HR_RE = re.compile(r"^(-{3,}|\*{3,}|_{3,})$")
_BULLET_RE = re.compile(r"^[-*+]\s+(.*)$")
_NUMBERED_RE = re.compile(r"^\d+[.)]\s+(.*)$")
_SEP_CELL_RE = re.compile(r"^:?-{2,}:?$")


def markdown_to_docx_bytes(markdown: str, title: str | None = None) -> bytes:
    from docx import Document
    from docx.shared import Pt

    doc = Document()
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10.5)

    if title:
        doc.core_properties.title = title

    lines = markdown.replace("\r\n", "\n").split("\n")
    i, n = 0, len(lines)
    while i < n:
        raw = lines[i]
        stripped = raw.strip()

        if not stripped:
            i += 1
            continue

        if stripped.startswith("```"):
            i += 1
            code: List[str] = []
            while i < n and not lines[i].strip().startswith("```"):
                code.append(lines[i])
                i += 1
            i += 1  # skip closing fence
            _add_code_block(doc, code)
            continue

        if "|" in stripped and _table_run_length(lines, i) >= 2:
            i = _consume_table(doc, lines, i)
            continue

        heading = _HEADING_RE.match(stripped)
        if heading:
            level = len(heading.group(1))
            _add_heading(doc, heading.group(2).strip(), level)
            i += 1
            continue

        if _HR_RE.match(stripped):
            i += 1
            continue

        bullet = _BULLET_RE.match(stripped)
        if bullet:
            para = doc.add_paragraph(style="List Bullet")
            _add_inline(para, bullet.group(1))
            i += 1
            continue

        numbered = _NUMBERED_RE.match(stripped)
        if numbered:
            para = doc.add_paragraph(style="List Number")
            _add_inline(para, numbered.group(1))
            i += 1
            continue

        para = doc.add_paragraph()
        _add_inline(para, stripped)
        i += 1

    bio = BytesIO()
    doc.save(bio)
    return bio.getvalue()


def _add_heading(doc, text: str, level: int) -> None:
    clean = re.sub(r"[*`]", "", text).strip()
    doc.add_heading(clean, level=min(max(level - 1, 0), 4))


def _add_code_block(doc, code_lines: List[str]) -> None:
    from docx.shared import Pt, RGBColor

    para = doc.add_paragraph()
    run = para.add_run("\n".join(code_lines))
    run.font.name = "Consolas"
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)


def _add_inline(paragraph, text: str) -> None:
    text = text.replace("\\|", "|")
    pos = 0
    for match in _INLINE_RE.finditer(text):
        if match.start() > pos:
            paragraph.add_run(text[pos:match.start()])
        token = match.group(0)
        if token.startswith("**"):
            run = paragraph.add_run(token[2:-2])
            run.bold = True
        else:
            run = paragraph.add_run(token[1:-1])
            run.font.name = "Consolas"
        pos = match.end()
    if pos < len(text):
        paragraph.add_run(text[pos:])


def _split_row(line: str) -> List[str]:
    body = line.strip()
    if body.startswith("|"):
        body = body[1:]
    if body.endswith("|"):
        body = body[:-1]
    body = body.replace("\\|", "\x00")
    cells = [c.strip().replace("\x00", "|") for c in body.split("|")]
    return cells


def _is_separator_row(cells: List[str]) -> bool:
    non_empty = [c for c in cells if c]
    return bool(non_empty) and all(_SEP_CELL_RE.match(c) for c in non_empty)


def _table_run_length(lines: List[str], start: int) -> int:
    count = 0
    j = start
    while j < len(lines) and "|" in lines[j] and lines[j].strip():
        count += 1
        j += 1
    return count


def _consume_table(doc, lines: List[str], start: int) -> int:
    j = start
    rows: List[List[str]] = []
    while j < len(lines) and "|" in lines[j] and lines[j].strip():
        cells = _split_row(lines[j])
        if not _is_separator_row(cells):
            rows.append(cells)
        j += 1

    if not rows:
        return j

    ncols = max(len(r) for r in rows)
    table = doc.add_table(rows=0, cols=ncols)
    try:
        table.style = "Table Grid"
    except KeyError:
        pass

    for idx, cells in enumerate(rows):
        cells = cells + [""] * (ncols - len(cells))
        row_cells = table.add_row().cells
        for col, value in enumerate(cells):
            cell = row_cells[col]
            cell.text = ""
            para = cell.paragraphs[0]
            _add_inline(para, value)
            if idx == 0:
                for run in para.runs:
                    run.bold = True
    return j
